# PDF Tool Web App using Flask (Python)
# Rx Advance PDF Tools: Compress PDF, Merge PDF, Convert PDF to Word, Image to PDF, View, Download & Delete Files

from flask import Flask, request, send_file, render_template_string, abort, redirect, url_for
import os
import subprocess
from PyPDF2 import PdfMerger
from datetime import datetime
import uuid
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import io
import pdfplumber
import pandas as pd
import pytesseract
import cv2
import numpy as np
from pdf2image import convert_from_path
from PIL import Image
from PyPDF2 import PdfReader, PdfWriter

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
SAVED_FOLDER = 'saved'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(SAVED_FOLDER, exist_ok=True)

HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Rx Advance PDF Tools</title>
  <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">
</head>
<body class="bg-gray-100">
  <div class="container mx-auto p-6">
    <header class="text-center mb-6">
      <h1 class="text-4xl font-bold text-blue-600">Rx Advance PDF Tools</h1>
      <p class="text-lg text-gray-700">Compress, Merge, Convert PDFs</p>
    </header>

    <div class="grid grid-cols-1 md:grid-cols-2 lg:grid-cols-3 gap-6">
      <div class="bg-white p-4 rounded-lg shadow-md text-center">
        <h2 class="text-xl font-semibold text-blue-500">Compress PDF</h2>
        <form action="/compress" method="post" enctype="multipart/form-data" onsubmit="showLoading(this)">
          <input type="file" name="pdf" accept="application/pdf" class="mt-2 block mx-auto" required>
          <button type="submit" class="mt-4 px-4 py-2 bg-blue-500 text-white rounded">Compress</button>
        </form>
      </div>

      <div class="bg-white p-4 rounded-lg shadow-md text-center">
  <h2 class="text-xl font-semibold text-blue-500">Merge PDFs</h2>
  <form action="/merge" method="post" enctype="multipart/form-data" onsubmit="return validateMultipleMerge(this)">
    <input type="file" name="pdfs" accept="application/pdf" class="mt-2 block mx-auto" multiple required>
    <p class="text-sm text-gray-500 mt-1">(You can upload 2 or more PDFs)</p>
    <button type="submit" class="mt-4 px-4 py-2 bg-blue-500 text-white rounded">Merge</button>
  </form>
</div>

      <div class="bg-white p-4 rounded-lg shadow-md text-center">
        <h2 class="text-xl font-semibold text-blue-500">PDF to Word</h2>
        <form action="/pdf_to_word" method="post" enctype="multipart/form-data" onsubmit="showLoading(this)">
          <input type="file" name="pdf" accept="application/pdf" class="mt-2 block mx-auto" required>
          <button type="submit" class="mt-4 px-4 py-2 bg-blue-500 text-white rounded">Convert</button>
        </form>
      </div>

      <div class="bg-white p-4 rounded-lg shadow-md text-center">
        <h2 class="text-xl font-semibold text-blue-500">Image to PDF</h2>
        <form action="/image_to_pdf" method="post" enctype="multipart/form-data" onsubmit="showLoading(this)">
          <input type="file" name="images" accept="image/*" class="mt-2 block mx-auto" multiple required>
          <button type="submit" class="mt-4 px-4 py-2 bg-blue-500 text-white rounded">Convert</button>
        </form>
      </div>
    
    
    <div class="bg-white p-4 rounded-lg shadow-md text-center">
  <h2 class="text-xl font-semibold text-blue-500">PDF to Excel</h2>
  <form action="/pdf_to_excel" method="post" enctype="multipart/form-data" onsubmit="showLoading(this)">
    <input type="file" name="pdf" accept="application/pdf" class="mt-2 block mx-auto" required>
    <button type="submit" class="mt-4 px-4 py-2 bg-blue-500 text-white rounded">Convert</button>
  </form>
</div>

<div class="bg-white p-4 rounded-lg shadow-md text-center">
  <h2 class="text-xl font-semibold text-blue-500">OCR Image to Word</h2>
  <form action="/ocr_image_to_docx" method="post" enctype="multipart/form-data" onsubmit="showLoading(this)">
    <input type="file" name="image" accept="image/*" class="mt-2 block mx-auto" required>
    <button type="submit" class="mt-4 px-4 py-2 bg-blue-500 text-white rounded">Convert</button>
  </form>
</div>
<div class="bg-white p-4 rounded-lg shadow-md text-center">
  <h2 class="text-xl font-semibold text-blue-500">JPG to PDF</h2>
  <form action="/jpg_to_pdf" method="post" enctype="multipart/form-data" onsubmit="showLoading(this)">
    <input type="file" name="images" accept=".jpg,.jpeg" class="mt-2 block mx-auto" multiple required>
    <button type="submit" class="mt-4 px-4 py-2 bg-blue-500 text-white rounded">Convert</button>
  </form>
</div>



<div class="bg-white p-4 rounded-lg shadow-md text-center">
  <h2 class="text-xl font-semibold text-blue-500">Excel to PDF</h2>
  <form action="/excel_to_pdf" method="post" enctype="multipart/form-data" onsubmit="showLoading(this)">
    <input type="file" name="excel" accept=".xls,.xlsx" class="mt-2 block mx-auto" required>
    <button type="submit" class="mt-4 px-4 py-2 bg-blue-500 text-white rounded">Convert</button>
  </form>
</div>


<div class="bg-white p-4 rounded-lg shadow-md text-center">
  <h2 class="text-xl font-semibold text-blue-500">Remove Pages from PDF</h2>
  <form action="/remove_pages" method="post" enctype="multipart/form-data" onsubmit="showLoading(this)">
    <input type="file" name="pdf" accept="application/pdf" class="mt-2 block mx-auto" required>
    <input type="text" name="pages" placeholder="Pages to remove (e.g., 2,4,6)" class="mt-2 block mx-auto border rounded p-1 w-3/4" required>
    <button type="submit" class="mt-4 px-4 py-2 bg-red-500 text-white rounded">Remove Pages</button>
  </form>
</div>

<div class="bg-white p-4 rounded-lg shadow-md text-center">
  <h2 class="text-xl font-semibold text-blue-500">Rotate Image</h2>
  <form action="/rotate_image" method="post" enctype="multipart/form-data" onsubmit="showLoading(this)">
    <input type="file" name="image" accept="image/*" class="mt-2 block mx-auto" required>
    <input type="number" name="angle" placeholder="Angle (e.g. 90, 180)" class="mt-2 block mx-auto border rounded p-1 w-3/4" required>
    <button type="submit" class="mt-4 px-4 py-2 bg-yellow-500 text-white rounded">Rotate</button>
  </form>
</div>


<div class="bg-white p-4 rounded-lg shadow-md text-center">
  <h2 class="text-xl font-semibold text-blue-500">Rotate PDF (Custom Per Page)</h2>
  <form action="/rotate_pdf" method="post" enctype="multipart/form-data" onsubmit="showLoading(this)">
    <input type="file" name="pdf" accept="application/pdf" class="mt-2 block mx-auto" required>
    <input type="text" name="rotations" placeholder="e.g. 1:90,2:180,5:270" class="mt-2 block mx-auto border rounded p-1 w-3/4" required>
    <button type="submit" class="mt-4 px-4 py-2 bg-purple-600 text-white rounded">Rotate Pages</button>
  </form>
</div>





</div>

    <div class="mt-8 text-center">
      <a href="/saved_files" class="text-blue-500 underline">📁 View Saved Files</a>
    </div>
  </div>

  <div id="spinner" class="hidden fixed top-0 left-0 w-full h-full bg-gray-700 bg-opacity-75 flex items-center justify-center z-50">
    <div class="text-white text-lg">Processing...</div>
  </div>

  <script>
    function showLoading(form) {
      document.getElementById('spinner').classList.remove('hidden');
      const btn = form.querySelector('button');
      btn.disabled = true;
      btn.innerText = 'Processing...';
    }

      function validateMultipleMerge(form) {
    const files = form.querySelector('input[type="file"]').files;
    if (files.length < 2) {
      alert("Please upload at least 2 PDF files to merge.");
      return false;
    }
    showLoading(form);
    return true;
  }
      showLoading(form);
      return true;
    }
  </script>
</body>
</html>
'''

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/saved_files')
def saved_files():
    files = os.listdir(SAVED_FOLDER)
    links = [
        f'<li class="mb-2">{f} '
        f'<a href="/download/{f}" class="text-green-600 underline">Download</a> '
        f'<a href="/delete/{f}" class="text-red-600 underline ml-4">Delete</a></li>'
        for f in files
    ]
    return '<h2 class="text-2xl text-center mt-4">Saved Files</h2><ul class="p-6">' + ''.join(links) + '</ul>'

@app.route('/download/<filename>')
def download_file(filename):
    path = os.path.join(SAVED_FOLDER, filename)
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    else:
        abort(404)

@app.route('/delete/<filename>')
def delete_file(filename):
    path = os.path.join(SAVED_FOLDER, filename)
    if os.path.exists(path):
        os.remove(path)
    return redirect(url_for('saved_files'))

@app.route('/compress', methods=['POST'])
def compress():
    file = request.files['pdf']
    if file:
        filename = f"compressed_{uuid.uuid4().hex}.pdf"
        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        output_path = os.path.join(SAVED_FOLDER, filename)
        file.save(input_path)
        compress_pdf(input_path, output_path)
        return redirect(url_for('download_file', filename=filename))

@app.route('/merge', methods=['POST'])
def merge():
    files = request.files.getlist('pdfs')
    if len(files) < 2:
        return "At least 2 PDF files are required to merge.", 400

    filename = f"merged_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(SAVED_FOLDER, filename)

    merger = PdfMerger()
    for file in files:
        file_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(file_path)
        merger.append(file_path)
    merger.write(output_path)
    merger.close()

    return redirect(url_for('download_file', filename=filename))

@app.route('/pdf_to_word', methods=['POST'])
def pdf_to_word():
      file = request.files['pdf']
      if file:
       filename = f"converted_{uuid.uuid4().hex}.docx"
      input_path = os.path.join(UPLOAD_FOLDER, file.filename)
      output_path = os.path.join(SAVED_FOLDER, filename)
      file.save(input_path)
      convert_pdf_to_word(input_path, output_path)
      return redirect(url_for('download_file', filename=filename))

@app.route('/image_to_pdf', methods=['POST'])
def image_to_pdf():
    files = request.files.getlist('images')
    images = []
    for file in files:
        img = Image.open(file.stream).convert('RGB')
        images.append(img)

    filename = f"imagepdf_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(SAVED_FOLDER, filename)

    if images:
        images[0].save(output_path, save_all=True, append_images=images[1:])
        return redirect(url_for('download_file', filename=filename))
    else:
        return "No valid images uploaded", 400



@app.route('/pdf_to_excel', methods=['POST'])
def pdf_to_excel():
        file = request.files['pdf']
        if file:
            filename = f"excel_{uuid.uuid4().hex}.xlsx"
            input_path = os.path.join(UPLOAD_FOLDER, file.filename)
            output_path = os.path.join(SAVED_FOLDER, filename)
            file.save(input_path)
            convert_pdf_to_excel(input_path, output_path)
            return redirect(url_for('download_file', filename=filename))
        else:
            return "No PDF file uploaded", 400

@app.route('/ocr_image_to_docx', methods=['POST'])
def ocr_image_to_docx_route():
        file = request.files['image']
        if file:
            filename = f"ocr_{uuid.uuid4().hex}.docx"
            output_path = os.path.join(SAVED_FOLDER, filename)
            ocr_image_to_docx(file.stream, output_path)
            return redirect(url_for('download_file', filename=filename))
            return "No image uploaded", 400

@app.route('/jpg_to_pdf', methods=['POST'])
def jpg_to_pdf():
    files = request.files.getlist('images')
    if not files:
        return "No images uploaded", 400

    filename = f"jpg_to_pdf_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(SAVED_FOLDER, filename)

    convert_jpgs_to_pdf(files, output_path)

    return redirect(url_for('download_file', filename=filename))



@app.route('/excel_to_pdf', methods=['POST'])
def excel_to_pdf():
    file = request.files['excel']
    if file and file.filename.endswith(('.xls', '.xlsx')):
        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        output_filename = f"excelpdf_{uuid.uuid4().hex}.pdf"
        output_path = os.path.join(SAVED_FOLDER, output_filename)

        file.save(input_path)
        convert_excel_to_pdf(input_path, output_path)

        return redirect(url_for('download_file', filename=output_filename))
    return "Please upload a valid Excel file", 400


@app.route('/remove_pages', methods=['POST'])
def remove_pages():
    file = request.files['pdf']
    pages_str = request.form.get('pages')  # e.g., "1,3,5"
    if file and pages_str:
        pages_to_remove = [int(p.strip()) - 1 for p in pages_str.split(',') if p.strip().isdigit()]

        input_path = os.path.join(UPLOAD_FOLDER, file.filename)
        output_filename = f"removed_pages_{uuid.uuid4().hex}.pdf"
        output_path = os.path.join(SAVED_FOLDER, output_filename)

        file.save(input_path)
        remove_pages_from_pdf(input_path, output_path, pages_to_remove)

        return redirect(url_for('download_file', filename=output_filename))
    return "Invalid input", 400


@app.route('/rotate_image', methods=['POST'])
def rotate_image_route():
    file = request.files.get('image')
    angle = request.form.get('angle', type=int)

    if not file or not angle:
        return "Please upload an image and specify a rotation angle.", 400

    filename = f"rotated_{uuid.uuid4().hex}.jpg"
    output_path = os.path.join(SAVED_FOLDER, filename)

    rotate_image(file, angle, output_path)

    return redirect(url_for('download_file', filename=filename))


@app.route('/rotate_pdf', methods=['POST'])
def rotate_pdf():
    file = request.files['pdf']
    rotation_str = request.form.get('rotations', '')  # e.g., "1:90, 2:180"

    if not file or not rotation_str:
        return "Please upload a PDF and specify page:angle pairs.", 400

    rotation_map = parse_rotation_map(rotation_str)
    if not rotation_map:
        return "Invalid rotation format.", 400

    input_path = os.path.join(UPLOAD_FOLDER, file.filename)
    output_filename = f"rotated_pdf_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(SAVED_FOLDER, output_filename)

    file.save(input_path)
    rotate_pdf_pages_per_page(input_path, output_path, rotation_map)

    return redirect(url_for('download_file', filename=output_filename))








def compress_pdf(input_path: str, output_path: str, dpi: int = 72, quality: int = 50):
    doc = fitz.open(input_path)
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            pil_image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
            pil_image = pil_image.resize(
                (int(pil_image.width * dpi / 72), int(pil_image.height * dpi / 72)),
                Image.ANTIALIAS
            )
            output_image = io.BytesIO()
            pil_image.save(output_image, format="JPEG", quality=quality)
            output_image.seek(0)
            doc.update_image(xref, output_image.read())
    doc.save(output_path)

def convert_pdf_to_word(pdf_path: str, word_path: str):
    pdf_document = fitz.open(pdf_path)
    doc = Document()
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text()
        doc.add_paragraph(text)
    doc.save(word_path)

def convert_pdf_to_excel(pdf_path: str, excel_path: str):
        all_tables = []

        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        df = pd.DataFrame(table)
                        all_tables.append(df)

        if all_tables:
            with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                for i, df in enumerate(all_tables):
                    df.to_excel(writer, sheet_name=f"Page_{i + 1}", index=False, header=False)
        else:
            # Create an empty Excel file with a message
            df = pd.DataFrame([["No tables found in the PDF."]])
            df.to_excel(excel_path, index=False, header=False)

def ocr_image_to_docx(image_stream, output_path):
                pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
                # Read image from stream
                img = Image.open(image_stream).convert('RGB')
                open_cv_image = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)

                # Run OCR
                text = pytesseract.image_to_string(open_cv_image)

                # Save to DOCX
                doc = Document()
                doc.add_paragraph(text)
                doc.save(output_path)

def convert_jpgs_to_pdf(files, output_path):
    images = []

    for file in files:
        img = Image.open(file.stream).convert('RGB')
        images.append(img)

    if images:
        # Save all images into one PDF
        images[0].save(output_path, save_all=True, append_images=images[1:])





def convert_excel_to_pdf(input_path, output_path):
    excel = win32.gencache.EnsureDispatch('Excel.Application')
    excel.Visible = False
    wb = excel.Workbooks.Open(input_path)

    try:
        wb.ExportAsFixedFormat(0, output_path)
    finally:
        wb.Close(False)
        excel.Quit()

def remove_pages_from_pdf(input_path, output_path, pages_to_remove):
    reader = PdfReader(input_path)
    writer = PdfWriter()

    for i in range(len(reader.pages)):
        if i not in pages_to_remove:
            writer.add_page(reader.pages[i])

    with open(output_path, "wb") as f:
        writer.write(f)


def rotate_image(image_file, angle, output_path):
    image = Image.open(image_file.stream).convert('RGB')
    rotated = image.rotate(angle, expand=True)
    rotated.save(output_path, format="JPEG")

def rotate_pdf_pages_per_page(input_path, output_path, rotation_map):
    reader = PdfReader(input_path)
    writer = PdfWriter()

    for i in range(len(reader.pages)):
        page = reader.pages[i]
        angle = rotation_map.get(i, 0)
        if angle:
            page.rotate(angle)
        writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)

def parse_rotation_map(rotation_str):
    rotation_map = {}
    for item in rotation_str.split(','):
        if ':' in item:
            page_str, angle_str = item.split(':')
            if page_str.strip().isdigit() and angle_str.strip().lstrip('-').isdigit():
                page = int(page_str.strip()) - 1  # user input is 1-based
                angle = int(angle_str.strip())
                if angle % 90 == 0:
                    rotation_map[page] = angle
    return rotation_map



if __name__ == '__main__':
    app.run(debug=True)
